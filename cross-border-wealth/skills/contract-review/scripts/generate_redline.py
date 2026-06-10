#!/usr/bin/env python3
"""generate_redline.py — Generate a redline document comparing original and revised text.

Usage:
    python3 generate_redline.py --original <original.txt> --revised <revised.txt> --output <redline.docx>
    python3 generate_redline.py --original-text "original text" --revised-text "revised text" --output <redline.docx>

Requires python-docx for DOCX output. Falls back to markdown diff if not installed.
"""

import sys
import os
import argparse
import difflib
from datetime import datetime


def check_docx_available():
    """Check if python-docx is available."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def generate_docx_redline(original_lines, revised_lines, output_path, metadata=None):
    """Generate a DOCX file with tracked-changes-style formatting."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

    doc = Document()
    metadata = metadata or {}

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Contract Redline Comparison', level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(26, 29, 39)

    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_after = Pt(12)
    meta_items = [
        ('Date', metadata.get('date', datetime.now().strftime('%Y-%m-%d'))),
        ('Original', metadata.get('original_file', 'Original Document')),
        ('Revised', metadata.get('revised_file', 'Revised Document')),
    ]
    for label, value in meta_items:
        run = meta_para.add_run(f'{label}: ')
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run = meta_para.add_run(f'{value}    ')
        run.font.size = Pt(9)

    # Legend
    legend = doc.add_paragraph()
    legend.paragraph_format.space_after = Pt(18)
    run = legend.add_run('Legend: ')
    run.bold = True
    run.font.size = Pt(9)
    run = legend.add_run('Deleted text')
    run.font.color.rgb = RGBColor(220, 38, 38)
    run.font.strike = True
    run.font.size = Pt(9)
    run = legend.add_run('  |  ')
    run.font.size = Pt(9)
    run = legend.add_run('Added text')
    run.font.color.rgb = RGBColor(37, 99, 235)
    run.underline = True
    run.font.size = Pt(9)

    doc.add_paragraph('').paragraph_format.space_after = Pt(6)

    # Generate diff
    differ = difflib.SequenceMatcher(None, original_lines, revised_lines)

    for tag, i1, i2, j1, j2 in differ.get_opcodes():
        if tag == 'equal':
            for line in original_lines[i1:i2]:
                text = line.strip()
                if text:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.space_after = Pt(4)

        elif tag == 'delete':
            for line in original_lines[i1:i2]:
                text = line.strip()
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.color.rgb = RGBColor(220, 38, 38)
                    run.font.strike = True
                    para.paragraph_format.space_after = Pt(4)

        elif tag == 'insert':
            for line in revised_lines[j1:j2]:
                text = line.strip()
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.color.rgb = RGBColor(37, 99, 235)
                    run.underline = True
                    para.paragraph_format.space_after = Pt(4)

        elif tag == 'replace':
            # Show deletion then insertion inline in the same paragraph
            for line in original_lines[i1:i2]:
                text = line.strip()
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.color.rgb = RGBColor(220, 38, 38)
                    run.font.strike = True
                    para.paragraph_format.space_after = Pt(2)

            for line in revised_lines[j1:j2]:
                text = line.strip()
                if text:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.font.color.rgb = RGBColor(37, 99, 235)
                    run.underline = True
                    para.paragraph_format.space_after = Pt(4)

    doc.save(output_path)
    print(f'Redline DOCX saved to: {output_path}')
    return True


def generate_markdown_redline(original_lines, revised_lines, output_path):
    """Generate a markdown diff as fallback when python-docx is unavailable."""
    differ = difflib.SequenceMatcher(None, original_lines, revised_lines)
    output = []

    output.append('# Contract Redline Comparison')
    output.append('')
    output.append(f'Date: {datetime.now().strftime("%Y-%m-%d")}')
    output.append('')
    output.append('Legend: ~~deleted text~~ | **added text**')
    output.append('')
    output.append('---')
    output.append('')

    for tag, i1, i2, j1, j2 in differ.get_opcodes():
        if tag == 'equal':
            for line in original_lines[i1:i2]:
                text = line.strip()
                if text:
                    output.append(text)
                    output.append('')

        elif tag == 'delete':
            for line in original_lines[i1:i2]:
                text = line.strip()
                if text:
                    output.append(f'~~{text}~~')
                    output.append('')

        elif tag == 'insert':
            for line in revised_lines[j1:j2]:
                text = line.strip()
                if text:
                    output.append(f'**{text}**')
                    output.append('')

        elif tag == 'replace':
            for line in original_lines[i1:i2]:
                text = line.strip()
                if text:
                    output.append(f'~~{text}~~')
            for line in revised_lines[j1:j2]:
                text = line.strip()
                if text:
                    output.append(f'**{text}**')
            output.append('')

    md_path = output_path.replace('.docx', '.md') if output_path.endswith('.docx') else output_path
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(output))

    print(f'Markdown redline saved to: {md_path}')
    print('Note: Install python-docx for Word format output: pip install python-docx')
    return True


def main():
    parser = argparse.ArgumentParser(description='Generate a redline comparison document')
    parser.add_argument('--original', help='Path to original text file')
    parser.add_argument('--revised', help='Path to revised text file')
    parser.add_argument('--original-text', help='Original text as string')
    parser.add_argument('--revised-text', help='Revised text as string')
    parser.add_argument('--output', required=True, help='Output file path')
    args = parser.parse_args()

    # Load original text
    if args.original:
        if not os.path.exists(args.original):
            print(f'Error: Original file not found: {args.original}')
            sys.exit(1)
        with open(args.original, 'r', encoding='utf-8') as f:
            original_lines = f.readlines()
    elif args.original_text:
        original_lines = args.original_text.splitlines(keepends=True)
    else:
        print('Error: Provide --original or --original-text')
        sys.exit(1)

    # Load revised text
    if args.revised:
        if not os.path.exists(args.revised):
            print(f'Error: Revised file not found: {args.revised}')
            sys.exit(1)
        with open(args.revised, 'r', encoding='utf-8') as f:
            revised_lines = f.readlines()
    elif args.revised_text:
        revised_lines = args.revised_text.splitlines(keepends=True)
    else:
        print('Error: Provide --revised or --revised-text')
        sys.exit(1)

    # Build metadata
    metadata = {
        'date': datetime.now().strftime('%Y-%m-%d'),
        'original_file': args.original or 'inline text',
        'revised_file': args.revised or 'inline text',
    }

    # Generate output
    if check_docx_available():
        generate_docx_redline(original_lines, revised_lines, args.output, metadata)
    else:
        generate_markdown_redline(original_lines, revised_lines, args.output)


if __name__ == '__main__':
    main()
